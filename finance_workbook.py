from docx import Document

# Create a new Word document for Betegna Finance App
doc = Document()
doc.add_heading('Finance Workbook – Betegna Finance App (Ethiopia)', level=1)

# Worksheet 1
doc.add_heading('Worksheet 1: Budget Planning Worksheet', level=2)
doc.add_paragraph('Monthly Business Budget Template (ETB)')

table1 = doc.add_table(rows=1, cols=4)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Planned Amount (ETB)'
hdr_cells[2].text = 'Actual Amount (ETB)'
hdr_cells[3].text = 'Notes'

rows1 = [
    ('Income', '', '', ''),
    ('Sales Revenue', '220,000', '210,000', 'SME app subscriptions, PO finance interest fees'),
    ('Other Income (e.g., grants, partnerships)', '60,000', '50,000', 'Fintech innovation grant & consulting services'),
    ('Total Income', '280,000', '260,000', ''),
    ('Expenses', '', '', ''),
    ('Rent', '20,000', '20,000', 'Office in Addis Ababa'),
    ('Utilities (Electricity, Water, Internet)', '10,000', '9,500', 'Regular utilities & SaaS hosting'),
    ('Raw Materials / Supplies', '6,000', '5,800', 'Office materials, server costs'),
    ('Employee Wages', '130,000', '125,000', 'Developers, finance analyst, admin staff'),
    ('Transportation', '10,000', '9,000', 'Client meetings and travel'),
    ('Marketing / Promotion', '25,000', '23,000', 'SME outreach campaigns and online ads'),
    ('Loan Repayments', '0', '0', 'None in initial stage'),
    ('Equipment Maintenance', '5,000', '4,500', 'Computer & server upkeep'),
    ('Miscellaneous', '9,000', '8,000', 'Unexpected operating costs'),
    ('Total Expenses', '215,000', '204,800', ''),
    ('Net Profit / Loss (Total Income - Expenses)', '65,000', '55,200', 'Healthy startup margin Month 1')
]

for row in rows1:
    cells = table1.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Worksheet 2
doc.add_page_break()
doc.add_heading('Worksheet 2: Cash Flow Projection Worksheet', level=2)
doc.add_paragraph('3-Month Cash Flow Forecast (ETB)')

table2 = doc.add_table(rows=1, cols=6)
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Month'
hdr2[1].text = 'Cash Inflow (ETB)'
hdr2[2].text = 'Cash Outflow (ETB)'
hdr2[3].text = 'Net Cash Flow'
hdr2[4].text = 'Beginning Cash Balance'
hdr2[5].text = 'Ending Cash Balance'

rows2 = [
    ('Month 1', '280,000', '215,000', '+65,000', '400,000', '465,000'),
    ('Month 2', '320,000', '260,000', '+60,000', '465,000', '525,000'),
    ('Month 3', '360,000', '300,000', '+60,000', '525,000', '585,000'),
    ('Total 3 Months', '960,000', '775,000', '+185,000', '-', '585,000 (Final Balance)')
]

for row in rows2:
    cells = table2.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_paragraph(
    'Notes:\n- Steady revenue growth expected as more SMEs adopt the platform.\n'
    '- Cash inflow includes purchase order financing repayments with interest.\n'
    '- Expenses increase due to marketing and server capacity upgrades.'
)

# Worksheet 3
doc.add_page_break()
doc.add_heading('Worksheet 3: Product Costing & Pricing Worksheet', level=2)
doc.add_paragraph('Product Cost Breakdown and Pricing')
doc.add_paragraph('Product/Service Name: Betegna Finance App (All-in-One SME Finance Platform)')

table3 = doc.add_table(rows=1, cols=3)
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Cost Type'
hdr3[1].text = 'Amount (ETB)'
hdr3[2].text = 'Details'

rows3 = [
    ('Direct Materials', '5,000', 'Cloud server hosting, API tools'),
    ('Direct Labor (Wages per unit)', '25,000', 'App development and support team'),
    ('Packaging', '2,000', 'UI/UX design and branding'),
    ('Transport / Delivery', '1,500', 'Customer demos and site visits'),
    ('Utilities per unit', '3,000', 'Power, internet, data storage'),
    ('Marketing Cost', '5,000', 'Digital ad per client acquisition'),
    ('Miscellaneous', '2,000', 'Contingency and testing'),
    ('Total Direct Cost', '43,500', '')
]

for row in rows3:
    cells = table3.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_paragraph('Pricing Calculation')
doc.add_paragraph('Markup Percentage (%): 35%')
doc.add_paragraph('Selling Price = Total Cost + Markup')
doc.add_paragraph('Total Cost = 43,500 ETB')
doc.add_paragraph('Markup (35%) = 15,225 ETB')
doc.add_paragraph('Selling Price = 58,725 ETB per SME client package')

table4 = doc.add_table(rows=1, cols=4)
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Unit'
hdr4[1].text = 'Selling Price (ETB)'
hdr4[2].text = 'Total Cost (ETB)'
hdr4[3].text = 'Profit per Unit (ETB)'

data4 = [('1', '58,725', '43,500', '15,225')]

for row in data4:
    cells = table4.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

file_path = 'Finance_Workbook_Betegna_Finance_App.docx'
doc.save(file_path)

file_path
